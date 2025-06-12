import os
from pathlib import Path

from config import Config

from chromadb import PersistentClient
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from PyPDF2 import PdfReader
import docx


def load_documents(directory: str):
    """Load text from PDF, DOCX, and TXT files inside directory."""
    docs = []
    for root, _, files in os.walk(directory):
        for fname in files:
            path = Path(root) / fname
            lower = fname.lower()
            if lower.endswith('.txt'):
                text = Path(path).read_text(encoding='utf-8', errors='ignore')
            elif lower.endswith('.pdf'):
                reader = PdfReader(str(path))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            elif lower.endswith('.docx'):
                d = docx.Document(str(path))
                text = "\n".join(p.text for p in d.paragraphs)
            else:
                continue
            docs.append(Document(page_content=text, metadata={'source': str(path)}))
    return docs


def build_vector_store(documents):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=Config.CHUNK_SIZE,
        chunk_overlap=Config.CHUNK_OVERLAP,
    )
    texts = []
    metadatas = []
    for doc in documents:
        for chunk in splitter.split_text(doc.page_content):
            texts.append(chunk)
            metadatas.append(doc.metadata)

    if not texts:
        print("No content to add to the vector store.")
        return

    embedder = SentenceTransformer('all-MiniLM-L6-v2')
    embeddings = embedder.encode(texts).tolist()

    client = PersistentClient(path=Config.CHROMA_PERSIST_DIRECTORY)
    collection = client.get_or_create_collection('documents')
    ids = [str(i) for i in range(len(embeddings))]
    collection.add(embeddings=embeddings, metadatas=metadatas, ids=ids)
    print(f"Indexed {len(embeddings)} chunks.")


def main():
    upload_dir = Path(Config.UPLOAD_FOLDER)
    upload_dir.mkdir(parents=True, exist_ok=True)
    Path(Config.CHROMA_PERSIST_DIRECTORY).mkdir(parents=True, exist_ok=True)

    documents = load_documents(str(upload_dir))
    if not documents:
        print(f"No supported documents found in {upload_dir}.")
        return

    build_vector_store(documents)


if __name__ == "__main__":
    main()
